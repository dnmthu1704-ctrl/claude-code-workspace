#!/usr/bin/env python3
"""
format_notes.py — Chuyển recap họp thô → .docx chuẩn 2 phần:
  1. Nội dung trao đổi (theo mảng, viết lại chuyên nghiệp)
  2. Next steps (công việc | PIC | Deadline)
"""

import argparse
import json
import os
import re
import sys
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.exit("[x] Thiếu thư viện: pip3 install python-docx")

HERE = os.path.dirname(os.path.abspath(__file__))
CONFIG_PATH = os.path.join(HERE, "config.json")


def load_config():
    cfg = {"anthropic_api_key": "", "company": "SEONGON"}
    if os.path.exists(CONFIG_PATH):
        with open(CONFIG_PATH, "r", encoding="utf-8") as f:
            cfg.update(json.load(f))
    return cfg


def clean_text(text):
    """Bỏ @mention, emoji trang trí, dấu thừa."""
    text = re.sub(r"@\S+", "", text)
    emoji_pattern = re.compile(
        "[\U0001F300-\U0001F9FF\U00002702-\U000027B0\U0001F1E0-\U0001F1FF]+", flags=re.UNICODE
    )
    text = emoji_pattern.sub("", text)
    text = re.sub(r"\s{2,}", " ", text).strip()
    return text


def parse_with_ai(notes, client, date, attendees, api_key):
    """Dùng Claude API để phân tích và viết lại recap."""
    try:
        import anthropic
    except ImportError:
        return None

    prompt = f"""Bạn là trợ lý chuyên format biên bản họp cho công ty SEO/Marketing.

Thông tin buổi họp:
- Client: {client}
- Ngày: {date}
- Tham dự: {attendees}

Recap thô:
{notes}

Hãy hệ thống hoá và viết lại thành 2 phần rõ ràng, chuyên nghiệp bằng tiếng Việt:

PHẦN 1: NỘI DUNG TRAO ĐỔI
- Phân nhóm theo mảng công việc (GEO, CRO, SEO, Wireframe...)
- Tóm tắt súc tích, rõ ràng từng mảng
- Highlight các quyết định quan trọng và lưu ý cần nhớ
- KHÔNG dùng @mention hay emoji trang trí

PHẦN 2: NEXT STEPS
Liệt kê dạng bảng với 3 cột: Công việc | PIC | Deadline
- Trích xuất đầy đủ tất cả các task từ recap
- PIC: ghi tên người thực hiện (nếu có)
- Deadline: ghi ngày cụ thể (nếu có), nếu không ghi "Chưa xác định"

Trả về JSON với cấu trúc:
{{
  "sections": [
    {{
      "title": "tên mảng",
      "points": ["nội dung 1", "nội dung 2"]
    }}
  ],
  "next_steps": [
    {{"task": "tên công việc", "pic": "tên người", "deadline": "ngày"}}
  ]
}}"""

    client_obj = anthropic.Anthropic(api_key=api_key)
    response = client_obj.messages.create(
        model="claude-haiku-4-5-20251001",
        max_tokens=2000,
        messages=[{"role": "user", "content": prompt}],
    )
    text = response.content[0].text.strip()
    json_match = re.search(r"\{.*\}", text, re.DOTALL)
    if json_match:
        return json.loads(json_match.group())
    return None


# Từ khoá báo hiệu phần "deadline" khi dòng không có ngày tháng dạng d/m
# (vd: "... cho đến khi hết điểm chuẩn"). Xếp theo thứ tự dài -> ngắn để
# regex khớp đúng cụm dài nhất trước.
_DEADLINE_KEYWORDS = [
    "cho đến khi", "đến khi", "tới khi", "trước khi",
    "cho đến", "trước ngày", "trong tuần", "trong tháng",
    "tới", "hết", "đến",
]
_DEADLINE_KEYWORD_PATTERN = "|".join(re.escape(k) for k in _DEADLINE_KEYWORDS)


def split_pic_deadline(suffix):
    """Tách PIC + deadline từ phần văn bản sau dấu '-' cuối cùng của 1 dòng next step.

    Không hardcode danh sách tên — nhận diện qua vị trí ngày tháng hoặc từ khoá
    deadline, phần còn lại (thường ngắn) được coi là PIC.
    """
    suffix = suffix.strip(" ,.-–")
    if not suffix:
        return "Chưa xác định", "Chưa xác định"

    date_match = re.search(r"\d{1,2}/\d{1,2}(?:/\d{2,4})?", suffix)
    if date_match:
        pic = suffix[: date_match.start()].strip(" ,.-–")
        return (pic or "Chưa xác định"), date_match.group()

    kw_match = re.search(_DEADLINE_KEYWORD_PATTERN, suffix, re.I)
    if kw_match:
        pic = suffix[: kw_match.start()].strip(" ,.-–")
        deadline = suffix[kw_match.start():].strip(" ,.-–")
        return (pic or "Chưa xác định"), (deadline or "Chưa xác định")

    return suffix, "Chưa xác định"


def looks_like_action_line(line):
    """Nhận diện 1 dòng có phải next step hay không, dựa trên phần sau dấu '-' cuối cùng.

    Dùng khi recap không có tiêu đề "Next step:" tường minh (trường hợp phổ biến
    với recap dạng liệt kê thẳng task, vd copy từ Slack/Zalo).
    """
    parts = re.split(r"[-–]", line)
    if len(parts) < 2:
        return False
    suffix = parts[-1].strip()
    if not suffix:
        return False
    has_date = bool(re.search(r"\d{1,2}/\d{1,2}", suffix))
    has_keyword = bool(re.search(_DEADLINE_KEYWORD_PATTERN, suffix, re.I))
    short_enough = len(suffix.split()) <= 5
    return has_date or has_keyword or short_enough


def split_task_and_pic(line):
    """Tách task (trước dấu '-' cuối) và PIC/deadline (sau dấu '-' cuối) của 1 dòng."""
    idx = max(line.rfind("-"), line.rfind("–"))
    task = line[:idx].strip(" ,.-–")
    pic, deadline = split_pic_deadline(line[idx + 1:])
    return task, pic, deadline


def parse_basic(notes):
    """Parser cơ bản khi không có API key."""
    sections = []
    next_steps = []

    section_keywords = {
        "GEO": ["geo", "generative", "prompt", "apac", "thị trường"],
        "CRO": ["cro", "chuyển đổi", "conversion", "crm", "đo lường"],
        "SEO / AIO": ["seo", "aio", "bài viết", "outline", "content", "cta"],
        "Wireframe / Web": ["wireframe", "giao diện", "web", "blog", "thiết kế"],
    }

    lines = [l.strip() for l in notes.split("\n") if l.strip()]
    # Nếu recap có tiêu đề "Next step:" tường minh, chỉ coi các dòng SAU tiêu đề
    # đó là next step (giữ hành vi cũ). Nếu không có, tự nhận diện từng dòng
    # bằng looks_like_action_line() — tránh việc cả recap "mất trắng" next steps
    # chỉ vì thiếu đúng cụm từ "next step".
    has_explicit_header = any(re.search(r"next\s*step", clean_text(l), re.I) for l in lines)

    current_section = "Chung"
    section_map = {current_section: []}

    in_next = False
    for line in lines:
        line_clean = clean_text(line)
        if not line_clean or len(line_clean) < 5:
            continue

        if re.search(r"next\s*step", line_clean, re.I):
            in_next = True
            continue

        is_next_step_line = in_next or (
            not has_explicit_header and looks_like_action_line(line_clean)
        )

        if is_next_step_line:
            task, pic, deadline = split_task_and_pic(line_clean)
            if task and len(task) > 5:
                next_steps.append({"task": task, "pic": pic, "deadline": deadline})
            continue

        matched = False
        for sec, kws in section_keywords.items():
            if any(kw in line_clean.lower() for kw in kws):
                if sec not in section_map:
                    section_map[sec] = []
                current_section = sec
                matched = True
                break

        section_map.setdefault(current_section, []).append(line_clean)

    for title, points in section_map.items():
        filtered = [p for p in points if len(p) > 10]
        if filtered:
            sections.append({"title": title, "points": filtered})

    return {"sections": sections, "next_steps": next_steps}


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def build_docx(data, client, date, attendees, company, out_path):
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    # Tiêu đề
    title = doc.add_heading(f"BIÊN BẢN HỌP — {client.upper()}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

    # Thông tin chung
    doc.add_paragraph(f"Ngày họp: {date}")
    doc.add_paragraph(f"Thành phần: {attendees}")
    doc.add_paragraph(f"Lập bởi: {company}")
    doc.add_paragraph("")

    # PHẦN 1: Nội dung trao đổi
    h1 = doc.add_heading("PHẦN 1: NỘI DUNG TRAO ĐỔI", level=2)
    h1.runs[0].font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

    for section in data.get("sections", []):
        if not section.get("points"):
            continue
        doc.add_heading(section["title"], level=3)
        for point in section["points"]:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(point)

    doc.add_paragraph("")

    # PHẦN 2: Next steps
    h2 = doc.add_heading("PHẦN 2: NEXT STEPS", level=2)
    h2.runs[0].font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

    next_steps = data.get("next_steps", [])
    if next_steps:
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"

        headers = ["Công việc", "PIC", "Deadline"]
        hdr = table.rows[0]
        for i, h in enumerate(headers):
            cell = hdr.cells[i]
            cell.text = h
            set_cell_bg(cell, "1A73E8")
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        for item in next_steps:
            row = table.add_row()
            row.cells[0].text = item.get("task", "")
            row.cells[1].text = item.get("pic", "")
            row.cells[2].text = item.get("deadline", "")
    else:
        doc.add_paragraph("(Không có next steps trong recap)")

    os.makedirs(os.path.dirname(out_path) if os.path.dirname(out_path) else ".", exist_ok=True)
    doc.save(out_path)

    return len(data.get("sections", [])), len(next_steps)


def main():
    ap = argparse.ArgumentParser(description="Format recap họp thô → Google Docs")
    ap.add_argument("--client", required=True)
    ap.add_argument("--date", required=True)
    ap.add_argument("--attendees", default="")
    ap.add_argument("--notes", default="")
    ap.add_argument("--notes-file", default="")
    ap.add_argument("--out", required=True)
    args = ap.parse_args()

    if args.notes_file:
        with open(args.notes_file, "r", encoding="utf-8") as f:
            notes = f.read()
    else:
        notes = args.notes

    if not notes.strip():
        sys.exit("[x] Không có nội dung recap. Dùng --notes hoặc --notes-file.")

    cfg = load_config()
    api_key = cfg.get("anthropic_api_key", "")

    print(f"[i] Client: {args.client} | Ngày: {args.date}")

    data = None
    if api_key:
        print("[i] Dùng Claude AI để phân tích...")
        try:
            data = parse_with_ai(notes, args.client, args.date, args.attendees, api_key)
        except Exception as e:
            print(f"[!] AI lỗi: {e} → dùng parser cơ bản")

    if not data:
        print("[i] Dùng parser cơ bản...")
        data = parse_basic(notes)

    print("[i] Xuất file .docx...")
    n_sections, n_steps = build_docx(
        data, args.client, args.date, args.attendees,
        cfg.get("company", "SEONGON"), args.out
    )
    print(f"[✓] Đã tạo: {args.out}")
    print(f"    → {n_sections} mảng nội dung, {n_steps} next steps")


if __name__ == "__main__":
    main()
