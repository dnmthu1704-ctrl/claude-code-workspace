#!/usr/bin/env python3
"""
monthly-report/generate_report.py
Đọc file Excel/CSV data SEO → tạo báo cáo tháng dạng .docx

Cách dùng:
    python3 generate_report.py --file data.xlsx --client "Vĩnh Tường" \
        --month "07-2026" --out outputs/bao-cao-seo-vinhthuong-07-2026.docx

    python3 generate_report.py --file data.xlsx --preview
"""

import argparse
import os
import sys
import csv
from datetime import datetime, timedelta
from collections import defaultdict

HERE = os.path.dirname(os.path.abspath(__file__))


def read_file(filepath, encoding="utf-8"):
    """Đọc file Excel hoặc CSV, trả về list of dicts."""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in (".xlsx", ".xls"):
        try:
            import openpyxl
        except ImportError:
            raise SystemExit("[x] Thiếu openpyxl. Chạy: pip3 install openpyxl")
        wb = openpyxl.load_workbook(filepath, read_only=True, data_only=True)
        ws = wb.active
        rows = list(ws.values)
        if not rows:
            raise SystemExit("[x] File Excel rỗng.")
        headers = [str(h).strip() if h is not None else f"col_{i}" for i, h in enumerate(rows[0])]
        return [dict(zip(headers, row)) for row in rows[1:] if any(v is not None for v in row)]
    elif ext == ".csv":
        rows = []
        with open(filepath, "r", encoding=encoding, errors="replace") as f:
            reader = csv.DictReader(f)
            for row in reader:
                rows.append(dict(row))
        return rows
    else:
        raise SystemExit(f"[x] Định dạng file không hỗ trợ: {ext}. Dùng .xlsx hoặc .csv")


def detect_columns(rows):
    """Tự động nhận diện các cột quan trọng từ tên cột."""
    if not rows:
        return {}
    cols = list(rows[0].keys())
    col_map = {}

    keyword_hints = ["keyword", "từ khóa", "tu khoa", "query", "search term", "term"]
    position_hints = ["position", "rank", "ranking", "vị trí", "vi tri", "pos"]
    traffic_hints = ["clicks", "traffic", "sessions", "click", "lượt click", "luot click"]
    impression_hints = ["impressions", "impression", "hiển thị", "hien thi"]
    url_hints = ["url", "page", "landing page", "trang"]

    for col in cols:
        col_lower = col.lower()
        if any(h in col_lower for h in keyword_hints) and "keyword" not in col_map:
            col_map["keyword"] = col
        if any(h in col_lower for h in position_hints) and "position" not in col_map:
            col_map["position"] = col
        if any(h in col_lower for h in traffic_hints) and "clicks" not in col_map:
            col_map["clicks"] = col
        if any(h in col_lower for h in impression_hints) and "impressions" not in col_map:
            col_map["impressions"] = col
        if any(h in col_lower for h in url_hints) and "url" not in col_map:
            col_map["url"] = col

    return col_map


def safe_float(val, default=0.0):
    try:
        return float(str(val).replace(",", "").strip())
    except (ValueError, TypeError):
        return default


def analyze_data(rows, col_map, top_n=10):
    """Phân tích dữ liệu SEO và trả về dict kết quả."""
    kw_col = col_map.get("keyword")
    pos_col = col_map.get("position")
    clicks_col = col_map.get("clicks")
    imp_col = col_map.get("impressions")

    total_keywords = len(rows)
    total_clicks = 0
    total_impressions = 0
    top10_keywords = []
    top3_keywords = []
    position_dist = defaultdict(int)

    for row in rows:
        pos = safe_float(row.get(pos_col, 0)) if pos_col else 0
        clicks = safe_float(row.get(clicks_col, 0)) if clicks_col else 0
        impressions = safe_float(row.get(imp_col, 0)) if imp_col else 0
        kw = str(row.get(kw_col, "")) if kw_col else ""

        total_clicks += clicks
        total_impressions += impressions

        if 0 < pos <= 3:
            position_dist["Top 3"] += 1
            top3_keywords.append({"keyword": kw, "position": pos, "clicks": clicks})
        elif 3 < pos <= 10:
            position_dist["Top 4-10"] += 1
        elif 10 < pos <= 20:
            position_dist["Top 11-20"] += 1
        else:
            position_dist["Ngoài Top 20"] += 1

        if 0 < pos <= top_n:
            top10_keywords.append({"keyword": kw, "position": pos, "clicks": clicks})

    top10_keywords = sorted(top10_keywords, key=lambda x: x["position"])[:top_n]
    ctr = (total_clicks / total_impressions * 100) if total_impressions > 0 else 0

    return {
        "total_keywords": total_keywords,
        "total_clicks": int(total_clicks),
        "total_impressions": int(total_impressions),
        "ctr": round(ctr, 2),
        "top10_keywords": top10_keywords,
        "top3_count": position_dist["Top 3"],
        "top10_count": position_dist["Top 3"] + position_dist["Top 4-10"],
        "position_dist": dict(position_dist),
    }


def write_docx(stats, client, month, out_path):
    """Xuất báo cáo SEO tháng ra file .docx."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise SystemExit("[x] Thiếu python-docx. Chạy: pip3 install python-docx")

    doc = Document()

    # Tiêu đề
    title = doc.add_heading(f"BÁO CÁO SEO THÁNG {month}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph(f"Client: {client}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(13)

    doc.add_paragraph()

    # I. Tổng quan
    doc.add_heading("I. TỔNG QUAN", level=1)
    overview_table = doc.add_table(rows=4, cols=2)
    overview_table.style = "Table Grid"
    overview_data = [
        ("Tổng số từ khóa theo dõi", f"{stats['total_keywords']:,}"),
        ("Từ khóa Top 10", f"{stats['top10_count']:,}"),
        ("Từ khóa Top 3", f"{stats['top3_count']:,}"),
        ("Tổng lượt click (tháng)", f"{stats['total_clicks']:,}"),
    ]
    for i, (label, value) in enumerate(overview_data):
        overview_table.rows[i].cells[0].text = label
        overview_table.rows[i].cells[1].text = value
        overview_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # II. Phân bổ từ khóa theo vị trí
    doc.add_heading("II. PHÂN BỔ TỪ KHÓA THEO VỊ TRÍ", level=1)
    dist_table = doc.add_table(rows=1, cols=2)
    dist_table.style = "Table Grid"
    dist_table.rows[0].cells[0].text = "Nhóm vị trí"
    dist_table.rows[0].cells[1].text = "Số từ khóa"
    dist_table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    dist_table.rows[0].cells[1].paragraphs[0].runs[0].bold = True

    for group, count in stats["position_dist"].items():
        row = dist_table.add_row()
        row.cells[0].text = group
        row.cells[1].text = str(count)

    doc.add_paragraph()

    # III. Top từ khóa
    doc.add_heading(f"III. TOP {len(stats['top10_keywords'])} TỪ KHÓA NỔI BẬT", level=1)
    if stats["top10_keywords"]:
        kw_table = doc.add_table(rows=1, cols=3)
        kw_table.style = "Table Grid"
        for i, h in enumerate(["Từ khóa", "Vị trí", "Lượt click"]):
            kw_table.rows[0].cells[i].text = h
            kw_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        for kw in stats["top10_keywords"]:
            row = kw_table.add_row()
            row.cells[0].text = str(kw["keyword"])
            row.cells[1].text = str(int(kw["position"])) if kw["position"] else "-"
            row.cells[2].text = str(int(kw["clicks"])) if kw["clicks"] else "-"
    else:
        doc.add_paragraph("Chưa có dữ liệu từ khóa Top 10.")

    doc.add_paragraph()

    # IV. Phân tích & đánh giá
    doc.add_heading("IV. PHÂN TÍCH & ĐÁNH GIÁ", level=1)
    doc.add_paragraph(
        f"Trong tháng {month}, website ghi nhận {stats['total_keywords']:,} từ khóa được theo dõi, "
        f"trong đó {stats['top10_count']:,} từ khóa đạt vị trí Top 10 và "
        f"{stats['top3_count']:,} từ khóa đạt Top 3. "
        f"Tổng lượt click từ organic search đạt {stats['total_clicks']:,} lượt."
    )

    doc.add_paragraph()

    # V. Kế hoạch tháng sau
    doc.add_heading("V. KẾ HOẠCH THÁNG SAU", level=1)
    plans = [
        "Tiếp tục tối ưu nội dung cho các từ khóa đang ở vị trí Top 11-20",
        "Xây dựng thêm backlink chất lượng cho các trang quan trọng",
        "Cải thiện tốc độ tải trang và Core Web Vitals",
        "Review và cập nhật nội dung các bài viết có traffic giảm",
    ]
    for plan in plans:
        doc.add_paragraph(plan, style="List Bullet")

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph(
        f"Báo cáo được tạo bởi Claude Code | SEONGON | {datetime.now().strftime('%d/%m/%Y')}"
    )
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)

    os.makedirs(os.path.dirname(out_path) if os.path.dirname(out_path) else ".", exist_ok=True)
    doc.save(out_path)
    return out_path


def main():
    ap = argparse.ArgumentParser(description="Tạo báo cáo SEO tháng từ file Excel/CSV")
    ap.add_argument("--file", required=True, help="File data SEO (.xlsx hoặc .csv)")
    ap.add_argument("--client", default="Client", help="Tên client")
    ap.add_argument("--month", default=datetime.now().strftime("%m-%Y"), help="Tháng báo cáo MM-YYYY")
    ap.add_argument("--top", type=int, default=10, help="Số từ khóa Top N (mặc định: 10)")
    ap.add_argument("--encoding", default="utf-8", help="Encoding file CSV")
    ap.add_argument("--out", default="bao-cao-seo.docx", help="File output .docx")
    ap.add_argument("--preview", action="store_true", help="Xem trước cấu trúc file, không tạo báo cáo")
    args = ap.parse_args()

    print(f"[i] Đọc file: {args.file}")
    rows = read_file(args.file, encoding=args.encoding)
    print(f"[i] Đọc được {len(rows)} dòng dữ liệu")

    col_map = detect_columns(rows)

    if args.preview:
        print("\n=== PREVIEW CẤU TRÚC FILE ===")
        print(f"Số dòng: {len(rows)}")
        print(f"Các cột: {list(rows[0].keys()) if rows else 'không có'}")
        print(f"\nCột được nhận diện:")
        for key, col in col_map.items():
            print(f"  {key:15} → {col}")
        if not col_map:
            print("  [!] Không nhận diện được cột tự động. Kiểm tra tên cột trong file.")
        print("\nDùng --out để tạo báo cáo.")
        return

    if not col_map:
        print("[!] Không nhận diện được cột. Script sẽ tạo báo cáo với dữ liệu cơ bản.")

    print(f"[i] Phân tích dữ liệu...")
    stats = analyze_data(rows, col_map, top_n=args.top)

    print(f"[i] Xuất báo cáo .docx...")
    out_path = write_docx(stats, args.client, args.month, args.out)

    print(f"[✓] Đã tạo: {out_path}")
    print(f"    → Tổng từ khóa: {stats['total_keywords']:,}")
    print(f"    → Top 10: {stats['top10_count']:,} | Top 3: {stats['top3_count']:,}")
    print(f"    → Tổng clicks: {stats['total_clicks']:,}")


if __name__ == "__main__":
    main()
